# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "LightBreaker用户手册.docx"

ACCENT = "2563EB"
ACCENT_DARK = "0D1528"
ACCENT_SOFT = "EAF1FF"
INK = "1F2937"
MUTED = "64748B"
LINE = "CBD5E1"
GREEN = "16A34A"
ORANGE = "EA580C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK if level == 1 else ACCENT)
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = add_paragraph(doc, text)
    paragraph.style = doc.styles["Body Text"]
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.18
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(INK)
        run.font.size = Pt(10.5)
    return paragraph


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        paragraph = add_paragraph(doc, "• " + item, style="Body Text")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Cm(0.2)
        for run in paragraph.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string(INK)


def add_numbered(doc: Document, items: list[str]):
    for index, item in enumerate(items, start=1):
        paragraph = add_paragraph(doc, f"{index}. {item}", style="Body Text")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Cm(0.2)
        for run in paragraph.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string(INK)


def add_note(doc: Document, title: str, text: str, fill: str = ACCENT_SOFT, color: str = ACCENT_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "BFD0EA")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(INK)
    set_run_font(run)
    doc.add_paragraph()
    return table


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True, color="FFFFFF", size=9)
        set_cell_shading(header_cells[idx], ACCENT_DARK)
        set_cell_border(header_cells[idx], ACCENT_DARK)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, size=9)
            set_cell_border(cells[idx])
            if idx == 0:
                set_cell_shading(cells[idx], "F8FAFC")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    for style_name in ("Body Text", "List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)

    for idx, size in ((1, 17), (2, 13), (3, 11)):
        style = styles[f"Heading {idx}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if idx == 1 else 8)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("LightBreaker")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    set_run_font(run)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("用户手册")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    set_run_font(run)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("蓝牙拳击手套 · 视觉拆盲盒 · 实时多人协作")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(run)

    doc.add_paragraph()
    add_note(
        doc,
        "适用版本",
        "适用于 LightBreaker Android MVP 及二期多人协作版本。内容覆盖单人训练、多人房间、蓝牙调试、画廊成长和常见问题处理。",
        fill="F8FAFC",
    )

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["文档版本", "V1.0"],
            ["生成日期", date.today().isoformat()],
            ["适用对象", "体验用户、现场运营人员、测试人员、项目维护人员"],
            ["推荐设备", "Android 手机、BOXING#Lxxxxxx 左手手套、BOXING#Rxxxxxx 右手手套"],
        ],
        widths=[3.5, 11.5],
    )
    doc.add_page_break()


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("LightBreaker 用户手册")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        set_run_font(run)


def build_manual() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. 产品简介")
    add_body(
        doc,
        "LightBreaker 是一款横屏拳击互动 APP。用户佩戴蓝牙拳击手套出拳，APP 根据手套上报的拳击次数和力度击碎屏幕上的瓷砖，逐步揭开隐藏在瓷砖下方的横屏图片。"
    )
    add_body(
        doc,
        "当前版本支持单人训练和云端实时多人协作。多人模式下，每台手机连接自己的手套，出拳事件上传到服务器排序后广播，所有参与者看到同一幅画面按相同顺序破砖。"
    )
    add_table(
        doc,
        ["模块", "用户能做什么"],
        [
            ["首页", "选择图片类型和难度，生成训练画作，扫描并连接手套，开始单人训练。"],
            ["多人", "创建房间或输入房间码加入，设置昵称，等待成员后开始实时协作训练。"],
            ["蓝牙", "扫描左右手手套，查看连接状态、电量、拳击次数、力度和原始数据包。"],
            ["游戏页", "横屏破砖，查看总出拳、揭开进度、连击、卡路里、左右手状态和多人贡献。"],
            ["结算页", "查看作品预览、拳数、击碎块数、最高连击、卡路里、XP 和保存状态。"],
            ["我的画廊", "查看已完成作品、生成提示词、完成时间和训练战绩摘要。"],
        ],
        widths=[3.0, 12.0],
    )

    add_heading(doc, "2. 使用前准备")
    add_bullets(
        doc,
        [
            "准备一台可运行 APP 的 Android 手机，并保持足够电量。",
            "开启手机蓝牙；Android 12 及以上系统需要允许蓝牙扫描和蓝牙连接权限。",
            "如需多人模式或云端图库，请保持手机网络可用。",
            "打开拳击手套电源，确认设备名称为 BOXING#L 加 6 位数字或字母的是左手，BOXING#R 加 6 位数字或字母的是右手。",
            "训练前留出安全挥拳空间，手机固定或握稳，儿童使用时建议成人在旁看护。",
        ],
    )
    add_note(
        doc,
        "安全提醒",
        "训练过程中请避免对人、玻璃、尖锐物或手机本体挥拳。若出现手腕不适、眩晕或疲劳，请立即停止训练。",
        fill="FFF7ED",
        color=ORANGE,
    )

    add_heading(doc, "3. 首次启动与权限")
    add_numbered(
        doc,
        [
            "安装并打开 LightBreaker。",
            "根据系统提示允许蓝牙相关权限；旧版 Android 可能会提示位置权限，这是系统对蓝牙扫描的要求。",
            "进入首页后查看顶部等级、XP 和最近连接设备。",
            "如看到默认画作信息，说明云端图片库可正常访问；也可以重新选择类型并生成画作。",
        ],
    )

    add_heading(doc, "4. 连接拳击手套")
    add_body(doc, "手套可以单只连接，也可以左右同时连接。单只手套适合快速体验，双手套适合完整训练和左右手统计。")
    add_numbered(
        doc,
        [
            "在首页点击“扫描手套”，或进入“蓝牙”页点击“扫描”。",
            "在设备列表中选择 BOXING#Lxxxxxx 或 BOXING#Rxxxxxx 设备。",
            "连接成功后，APP 会自动发送开启陀螺仪指令，并开始接收设备上报的拳击次数和力度。",
            "训练结束、退出训练或断开连接前，APP 会发送关闭陀螺仪指令。",
        ],
    )
    add_table(
        doc,
        ["状态", "含义", "建议操作"],
        [
            ["空闲", "尚未扫描或未连接设备。", "打开手套电源后点击扫描。"],
            ["扫描中", "正在查找 BOXING#Lxxxxxx / BOXING#Rxxxxxx 设备。", "等待设备出现，保持手套靠近手机。"],
            ["连接中", "已选择设备，正在建立蓝牙连接。", "不要关闭手套或切出 APP。"],
            ["就绪", "已连接并可接收拳击数据。", "可以开始训练。"],
            ["异常/断开", "连接失败或训练中断开。", "重新扫描连接；必要时重启手套。"],
        ],
        widths=[2.5, 6.2, 6.3],
    )

    add_heading(doc, "5. 单人训练")
    add_numbered(
        doc,
        [
            "进入“首页”。",
            "选择图片类型：自然风光、名画再现、城市建筑、抽象艺术、萌宠或科幻。",
            "选择难度：简单、标准或挑战。",
            "点击“生成画作”，等待当前画作信息刷新。",
            "确认至少一只手套处于就绪状态，然后点击“开始训练”。",
            "手机会进入横屏训练页，挥拳击碎瓷砖；如无设备，可用“模拟左拳/模拟右拳”做功能测试。",
            "完成 100% 揭开后自动结算，也可以点击“结算”提前结束。",
        ],
    )
    add_table(
        doc,
        ["难度", "瓷砖数量", "默认体验", "特点"],
        [
            ["简单", "约 150 块", "儿童、首次体验", "节奏轻，完成速度快。"],
            ["标准", "约 300 块", "默认 60 秒训练", "适合普通单人训练和现场体验。"],
            ["挑战", "约 500 块", "高强度训练", "包含锁定瓷砖、宝箱瓷砖和更高完成压力。"],
        ],
        widths=[2.4, 2.8, 4.0, 5.8],
    )
    add_note(
        doc,
        "画面反馈",
        "揭开进度达到 50% 时会出现流光反馈，80% 时会有粒子高亮，100% 时完整作品高亮展示。强力拳击会更容易击碎高硬度瓷砖或触发小范围震落。",
    )

    add_heading(doc, "6. 多人实时协作")
    add_body(
        doc,
        "多人模式使用匿名房间码，不需要账号登录。默认最多 4 人加入同一房间。所有玩家都可以自由击打整张图，不做区域分工，系统按玩家颜色和统计数据记录贡献。"
    )
    add_heading(doc, "6.1 房主创建房间", level=2)
    add_numbered(
        doc,
        [
            "进入“多人”页，输入昵称。",
            "选择图片类型和难度。",
            "点击“创建房间”。APP 会生成 6 位房间码。",
            "将房间码告诉其他成员。",
            "成员加入后，房主点击“开始协作”。",
        ],
    )
    add_heading(doc, "6.2 成员加入房间", level=2)
    add_numbered(
        doc,
        [
            "进入“多人”页，输入昵称。",
            "在房间码输入框填写房主提供的 6 位房间码。",
            "点击“加入房间”。",
            "等待房主开始训练；开始后手机会自动进入横屏训练页。",
        ],
    )
    add_heading(doc, "6.3 多人训练规则", level=2)
    add_bullets(
        doc,
        [
            "每台手机连接自己的左/右手套，出拳先上传到云端服务器。",
            "服务器为每次出拳分配递增序号 seq，再广播给房间内所有成员。",
            "所有客户端按相同 seq 顺序重放命中事件，因此画面进度保持一致。",
            "贡献统计包含出拳数、破砖贡献、最高连击、卡路里和 MVP 候选。",
            "断线后重新连接房间，可接收历史事件并尽量恢复到当前画面状态。",
        ],
    )

    add_heading(doc, "7. 宝箱、成长与奖励")
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["宝箱瓷砖", "占比约 3%-5%，击碎后随机触发奖励。"],
            ["范围震落", "对附近瓷砖造成额外破坏，加快揭图。"],
            ["双倍 XP", "结算时提高本次经验收益。"],
            ["快速破壁", "短时间内提升破砖效率。"],
            ["等级成长", "完成训练获得 XP，累积后提升破壁者等级。"],
            ["成就", "多人协作、宝箱发现、完整揭图等行为会逐步解锁成就。"],
        ],
        widths=[3.2, 11.8],
    )

    add_heading(doc, "8. 结算与我的画廊")
    add_body(doc, "训练结束后进入结算页。结算页会显示本轮作品、拳数、已击碎块数、最高连击、卡路里、XP、是否保存到画廊等信息。")
    add_bullets(
        doc,
        [
            "作品完整揭开 100% 后，会自动进入“我的画廊”。",
            "未完成的训练仍会保存训练记录，但不会作为完整作品进入画廊。",
            "多人完成作品会在摘要中显示房间信息、团队成员和贡献统计。",
            "画廊条目包含作品标题、图片类型、完成时间、出拳数、最高连击、卡路里和提示词。",
        ],
    )

    add_heading(doc, "9. 蓝牙调试页")
    add_body(doc, "蓝牙页用于现场测试和设备排障。普通用户不需要频繁进入，但在连接异常时非常有用。")
    add_table(
        doc,
        ["显示项", "用途"],
        [
            ["设备列表", "显示扫描到的左右手手套、RSSI 信号强度和设备名称。"],
            ["连接状态", "查看左/右手是否已连接、是否就绪、电量、拳击次数和力度。"],
            ["原始通知包", "显示 D5 5D 03 开头的数据包，供测试人员确认协议上报是否正常。"],
            ["断开全部", "结束调试或设备异常时，可断开所有蓝牙连接后重新扫描。"],
        ],
        widths=[3.5, 11.5],
    )

    add_heading(doc, "10. 常见问题")
    add_table(
        doc,
        ["问题", "可能原因", "处理方法"],
        [
            ["扫描不到手套", "手套未开机、距离过远、权限未允许。", "开启手套并靠近手机；确认蓝牙权限已允许；重新点击扫描。"],
            ["只连接到一只手套", "另一只手套未开机或未进入广播状态。", "检查设备名称是否符合 BOXING#L/R 加 6 位数字或字母；重启未出现的手套。"],
            ["连接成功但挥拳无反应", "设备未上报拳击次数或连接状态未就绪。", "进入蓝牙页查看是否有 D5 5D 03 数据包；断开后重新连接。"],
            ["多人房间无法加入", "房间码输入错误、房间已满、房间过期或网络不可用。", "确认 6 位房间码；检查网络；让房主重新创建房间。"],
            ["多人画面不同步", "网络延迟或短暂断线。", "保持网络稳定；离开后重新加入房间恢复事件历史。"],
            ["图片加载慢", "网络不稳定或云端图库响应慢。", "稍等片刻或切换网络后重新生成画作。"],
            ["结算后画廊没有作品", "本轮未达到 100% 完整揭开。", "完成整幅作品后再查看画廊。"],
        ],
        widths=[3.6, 5.4, 6.0],
    )

    add_heading(doc, "11. 运营与体验建议")
    add_bullets(
        doc,
        [
            "现场体验优先使用“标准 300”难度，节奏和完成概率更均衡。",
            "儿童或首次体验建议使用“简单 150”难度，并开启模拟拳测试流程。",
            "多人活动中建议由工作人员作为房主创建房间，确认所有成员就绪后再开始。",
            "每轮结束后引导用户查看画廊与等级成长，形成收藏和复玩动机。",
        ],
    )

    add_heading(doc, "12. 技术联调信息")
    add_table(
        doc,
        ["项目", "地址或说明"],
        [
            ["多人 REST", "http://152.136.62.157/lightbreaker/api/"],
            ["多人 WebSocket", "ws://152.136.62.157/lightbreaker/ws/"],
            ["图片目录", "http://152.136.62.157/lightbreaker/images/"],
            ["图片清单", "http://152.136.62.157/lightbreaker/images/manifest.json"],
            ["支持图片类型", "自然风光、名画再现、城市建筑、抽象艺术、萌宠、科幻"],
            ["蓝牙设备名", "左手 BOXING#L + 6 位数字或字母；右手 BOXING#R + 6 位数字或字母"],
        ],
        widths=[3.5, 11.5],
    )
    add_note(
        doc,
        "隐私说明",
        "二期 MVP 不需要账号登录。多人模式使用本地 installId、昵称和房间令牌识别参与者；用户昵称建议不要填写身份证号、手机号等敏感信息。",
        fill="F8FAFC",
    )

    add_footer(doc)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_manual()
